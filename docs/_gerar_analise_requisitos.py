"""Gera o DOCX de análise de requisitos — uso pontual."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

out_dir = Path(__file__).resolve().parent
path = out_dir / "Analise_Requisitos_Integracao_Stone_Tasy_Cotolengo.docx"

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_styled(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def add_para(text, bold=False, size=11, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=RGBColor(255, 255, 255))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()
    return table


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ANÁLISE DE REQUISITOS")
set_run_font(r, size=20, bold=True, color=RGBColor(31, 78, 121))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Integração Stone → Tasy\nHospital Pequeno Cotolengo")
set_run_font(r, size=14, bold=True)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Documento de apoio para elaboração da Análise de Requisitos formal\n"
    "Cliente: Pequeno Cotolengo  |  Fornecedor: GHR Tech  |  Agosto/2026"
)
set_run_font(r, size=10, color=RGBColor(89, 89, 89))
doc.add_paragraph()

add_heading_styled("1. Objetivo do projeto (visão simplificada)", 1)
add_para(
    "Eliminar o trabalho manual de conciliação das vendas das maquininhas Stone "
    "(cartão e PIX), automatizando a coleta dos extratos e a inserção das transações "
    "no sistema Tasy, com rastreabilidade, controle de erros e acompanhamento por portal."
)
add_para("Em uma frase:", bold=True)
add_para(
    "O que antes era baixar arquivo na Stone, digitar/conferir no Tasy e validar na mão "
    "passa a ser um fluxo automático diário: Stone → integração → Tasy, com painel de acompanhamento."
)

add_heading_styled("2. Escopo funcional", 1)
add_bullet("Extração automática de transações de cartão (crédito/débito/pré-pago) da API Stone.")
add_bullet("Recebimento de extratos PIX via solicitação + webhook oficial da Stone.")
add_bullet("Inserção no Tasy seguindo a regra de negócio do hospital (Caixa → Dia → Transação).")
add_bullet("Cadastro de maquininhas por setor/caixa (ex.: Cantina, Roupas, Mix, Telemarketing).")
add_bullet("Portal de controle: login, integrações, erros, filas, reprocessamento e monitoramento.")
add_bullet("Idempotência (não duplicar a mesma venda) e retry automático em falhas temporárias.")
add_para("Fora de escopo imediato (pode ser tratado em fase 2):", bold=True)
add_bullet("Substituição total de outras conciliações não-Stone.")
add_bullet("Alterações profundas de parametrização financeira no Tasy sem alinhamento do hospital.")

add_heading_styled("3. Situação atual — fluxo manual (como o cliente fazia)", 1)
add_para(
    "Antes da integração, o processo de conciliação dependia de pessoas e de várias etapas "
    "repetitivas. O fluxo típico era:"
)

add_heading_styled("3.1 Passo a passo manual", 2)
for s in [
    "Acessar o portal/site de conciliação da Stone com login institucional.",
    "Selecionar a data de movimento (em geral o dia anterior — D-1) e o estabelecimento/maquininha.",
    "Baixar o arquivo de extrato (Excel/CSV/XML conforme a tela disponível).",
    "Abrir o arquivo e conferir valores, bandeiras, tipos (crédito/débito/PIX) e terminais.",
    "Acessar o Tasy (módulo financeiro / tesouraria / cartões).",
    "Localizar o caixa do setor correspondente à maquininha (Cantina, Roupas, Mix, etc.).",
    "Abrir ou criar o movimento do dia (caixa diário).",
    "Lançar manualmente cada transação (ou em lote, quando possível) com valor, bandeira, autorização, parcelas e observações.",
    "Validar totais do arquivo Stone x totais lançados no Tasy (somatório, divergências, estornos).",
    "Corrigir erros de digitação, maquininha errada, caixa errado ou lançamento duplicado.",
    "Arquivar/guardar o Excel baixado para auditoria e eventual reconsulta.",
    "Repetir o processo para outros setores/maquininhas e, quando aplicável, para PIX em fluxo separado.",
]:
    add_numbered(s)

add_heading_styled("3.2 Diagrama do fluxo manual", 2)
add_para(
    "Usuário → Portal Stone → Download Excel → Conferência manual → Login Tasy → "
    "Seleção de caixa/setor → Digitação das vendas → Validação de totais → Correções → Arquivo guardado"
)
p = doc.add_paragraph()
run = p.add_run(
    "[1] Baixar extrato Stone (Excel)\n"
    "        ↓\n"
    "[2] Conferir linhas / valores / maquininha\n"
    "        ↓\n"
    "[3] Abrir Tasy e escolher caixa do setor\n"
    "        ↓\n"
    "[4] Lançar cartões (e PIX, se manual)\n"
    "        ↓\n"
    "[5] Validar totais e corrigir divergências\n"
    "        ↓\n"
    "[6] Encerrar conciliação do dia"
)
set_run_font(run, size=10)
p.paragraph_format.left_indent = Inches(0.3)

add_heading_styled("3.3 Dores e riscos do processo manual", 2)
add_table(
    ["Dor / risco", "Impacto"],
    [
        ["Tempo operacional alto", "Equipe gasta horas por dia/semana em download e digitação"],
        ["Erro humano", "Valor, bandeira, caixa ou serial incorretos"],
        ["Atraso na conciliação", "Fechamento financeiro depende da disponibilidade da pessoa"],
        ["Retrabalho", "Divergência só aparece na validação final"],
        ["Baixa rastreabilidade", "Difícil saber o que já foi lançado e o que falhou"],
        ["Dependência de arquivo", "Excel local pode se perder ou ficar desatualizado"],
        ["PIX separado", "Fluxo diferente do cartão aumenta complexidade"],
    ],
)

add_heading_styled("4. Situação proposta — o que o projeto entrega", 1)
add_para(
    "A solução automatiza a ponta a ponta: coleta na Stone, transporte seguro por fila, "
    "aplicação das regras de negócio e gravação no Tasy, com portal para acompanhar e reprocessar."
)

add_heading_styled("4.1 Objetivo técnico simplificado", 2)
add_bullet("Substituir download + digitação por integração via API/Webhook Stone.")
add_bullet("Garantir que cada venda (id Stone) entre uma única vez no Tasy.")
add_bullet("Separar cartão e PIX em fluxos próprios, sem misturar regras.")
add_bullet("Dar visibilidade ao hospital (portal) sem depender de planilha.")

add_heading_styled("4.2 Componentes da solução", 2)
add_table(
    ["Componente", "Função"],
    [
        ["stone-extracao", "Busca extratos na Stone (cartão) e recebe webhook PIX; publica na fila"],
        ["RabbitMQ", "Fila de mensagens entre extração e inserção (com retry/DLQ)"],
        ["tasy-insercao", "Aplica regras e insere no Oracle Tasy; grava status no Postgres"],
        ["Postgres (staging)", "Cadastros (maquininhas/caixas), status e auditoria da integração"],
        ["portal-controle", "Tela web para login, acompanhar integrações, erros, filas e cadastros"],
    ],
)

add_heading_styled("4.3 Fluxo automatizado — Cartão (rotina diária D-1)", 2)
for s in [
    "No horário agendado (ex.: madrugada/manhã), o serviço solicita o extrato do dia anterior na API Stone.",
    "O arquivo/retorno é parseado e cada transação vira uma mensagem na fila de cartão.",
    "O consumer lê a mensagem, identifica a maquininha → caixa/setor e a bandeira/tipo no mapeamento.",
    "Insere no Tasy na sequência Caixa → Dia → Transação (espelhando a regra já usada no hospital).",
    "Atualiza status no staging (integrado, erro, sem tesouraria, etc.).",
    "Em falha temporária, reenvia automaticamente (retry); se esgotar, preserva na DLQ para análise.",
    "O usuário acompanha no portal (totais, erros, reprocessamento) sem reabrir Excel.",
]:
    add_numbered(s)
add_para("Stone API (cartão D-1) → stone-extracao → Fila cartão → tasy-insercao → Tasy + Portal")

add_heading_styled("4.4 Fluxo automatizado — PIX", 2)
for s in [
    "O serviço solicita o extrato PIX da data desejada na API Stone (após a janela permitida pela Stone).",
    "A Stone processa e notifica o webhook HTTPS público do hospital com o link do arquivo.",
    "A integração baixa o CSV, parseia e publica na fila de PIX.",
    "O consumer aplica as regras (PIX tratado conforme regra de débito no Tasy) e grava o movimento.",
    "Status e erros ficam disponíveis no mesmo portal de controle.",
]:
    add_numbered(s)
add_para(
    "Stone (request PIX) → Webhook HTTPS → stone-extracao → Fila PIX → tasy-insercao → Tasy + Portal"
)

add_heading_styled("4.5 Diagrama do fluxo novo (visão única)", 2)
p = doc.add_paragraph()
run = p.add_run(
    "ANTES (manual)\n"
    "  Pessoa → Excel Stone → Digitação Tasy → Validação manual\n\n"
    "DEPOIS (automatizado)\n"
    "  Stone (API/Webhook)\n"
    "        ↓\n"
    "  Extração (stone-extracao)\n"
    "        ↓\n"
    "  Fila RabbitMQ (cartão / PIX)\n"
    "        ↓\n"
    "  Inserção (tasy-insercao) → Oracle Tasy\n"
    "        ↓\n"
    "  Portal de controle (acompanhamento / erros / reprocessar)"
)
set_run_font(run, size=10)
p.paragraph_format.left_indent = Inches(0.2)

add_heading_styled("5. Comparativo: antes x depois", 1)
add_table(
    ["Aspecto", "Antes (manual)", "Depois (projeto)"],
    [
        ["Origem dos dados", "Download Excel no site Stone", "API Stone + webhook PIX"],
        ["Lançamento no Tasy", "Digitação / importação manual", "Inserção automática via integração"],
        ["Frequência", "Quando a pessoa executa", "Rotina diária (D-1) + sob demanda"],
        ["Validação", "Conferência linha a linha", "Status por transação + painel de erros"],
        ["Duplicidade", "Risco de lançar 2x", "Idempotência por id Stone"],
        ["Falhas", "Retrabalho humano", "Retry automático + DLQ + reprocessar no portal"],
        ["Maquininha × setor", "Conhecimento da equipe", "Cadastro maquininha → caixa no staging"],
        ["PIX", "Fluxo à parte / manual", "Webhook oficial integrado à fila PIX"],
        ["Auditoria", "Arquivos Excel locais", "Histórico no staging + logs + portal"],
        ["Acesso da equipe", "Várias ferramentas", "Portal único de controle"],
    ],
)

add_heading_styled("6. Benefícios esperados", 1)
add_bullet("Redução drástica do tempo operacional de conciliação.")
add_bullet("Menos erro de digitação e de caixa/setor incorreto.")
add_bullet("Fechamento mais previsível (rotina D-1 automática).")
add_bullet("Visibilidade em tempo quase real do que integrou / falhou.")
add_bullet("Base para auditoria e suporte (status, observação, reprocessamento).")
add_bullet("Padronização cartão e PIX sob o mesmo modelo operacional.")

add_heading_styled("7. Requisitos de alto nível (para a analista detalhar)", 1)
add_heading_styled("7.1 Requisitos funcionais", 2)
add_table(
    ["ID", "Requisito"],
    [
        ("RF01", "Extrair conciliação de cartão da Stone para uma data (manual ou D-1)."),
        ("RF02", "Agendar execução diária automática do cartão D-1."),
        ("RF03", "Solicitar extrato PIX e receber notificação via webhook HTTPS."),
        ("RF04", "Publicar transações em filas distintas (cartão e PIX)."),
        ("RF05", "Inserir no Tasy respeitando Caixa → Dia → Transação."),
        ("RF06", "Mapear serial da maquininha para caixa/setor e transação financeira."),
        ("RF07", "Mapear bandeira/tipo Stone para códigos de cartão do Tasy."),
        ("RF08", "Garantir idempotência por identificador Stone (não duplicar)."),
        ("RF09", "Tratar falhas com retry e fila de dead-letter (DLQ)."),
        ("RF10", "Disponibilizar portal com login, listagem, filtros, erros e reprocessamento."),
        ("RF11", "Permitir cadastro/consulta de maquininhas e acompanhamento de filas."),
        ("RF12", "Tratar cenário de maquininha sem cadastro (ex.: inserção parcial / status específico)."),
    ],
)

add_heading_styled("7.2 Requisitos não funcionais", 2)
add_table(
    ["ID", "Requisito"],
    [
        ["RNF01", "Disponibilidade dos serviços em VM Windows do hospital (início automático)."],
        ["RNF02", "Webhook PIX acessível pela internet em HTTPS (DNS + certificado + proxy)."],
        ["RNF03", "Segurança de credenciais via variáveis de ambiente (.env), sem commit de segredos."],
        ["RNF04", "Logs e status auditáveis por transação."],
        ["RNF05", "Separação clara entre ambientes/homologação e produção."],
        ["RNF06", "Performance adequada ao volume diário de maquininhas do hospital."],
    ],
)

add_heading_styled("7.3 Dependências / premissas", 2)
add_bullet("Credenciais Stone (token cliente) e merchant IDs (StoneCode cartão / CNPJ PIX).")
add_bullet("Acesso Oracle Tasy (homolog/produção) e Postgres de staging.")
add_bullet("Cadastro das maquininhas com serial × caixa × código de transação financeira.")
add_bullet("Infraestrutura de rede: RabbitMQ, serviços Windows, e para PIX HTTPS público.")
add_bullet("Alinhamento dos códigos de bandeira/cartão no Tasy com o mapeamento da integração.")

add_heading_styled("8. Atores e responsabilidades", 1)
add_table(
    ["Ator", "Responsabilidade"],
    [
        ["Tesouraria / Financeiro", "Acompanhar portal, tratar erros de negócio, validar fechamento"],
        ["TI Hospital", "VM, DNS/HTTPS, firewall, proxy do subdomínio, Postgres/Rabbit"],
        ["GHR Tech", "Desenvolvimento, deploy, suporte à homologação e ajustes de regra"],
        ["Stone", "API de conciliação, webhook PIX e disponibilidade dos extratos"],
    ],
)

add_heading_styled("9. Critérios de aceite sugeridos", 1)
add_numbered("Cartão D-1 integrado automaticamente sem digitação manual no Tasy.")
add_numbered("Transação reprocessada não gera duplicidade no Tasy.")
add_numbered("Erros aparecem no portal com motivo compreensível para a operação.")
add_numbered("Maquininhas da lista do hospital direcionam para o caixa/setor correto.")
add_numbered("PIX: após HTTPS público, webhook registrado e extrato processado até o Tasy.")
add_numbered("Serviços sobem sozinhos após reinício da VM.")

add_heading_styled("10. Próximos passos para a analista", 1)
add_bullet("Detalhar casos de uso (UC) a partir dos RF01–RF12.")
add_bullet("Validar com o cliente o fluxo manual descrito na seção 3 (ajustar se houver variação local).")
add_bullet("Levantar volume médio diário de transações e setores críticos.")
add_bullet("Confirmar regras de PIX, pré-pago, parcelado e estorno com o financeiro.")
add_bullet("Incluir requisitos de infraestrutura (DNS stone.pequenocotolengo.org.br + HTTPS).")
add_bullet("Definir matriz de responsabilidades RACI e cronograma de homologação.")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    "Este documento é um insumo técnico-funcional para a analista formalizar a Análise de Requisitos "
    "no padrão interno do hospital/GHR. Pode ser expandido com wireframes do portal, dicionário de dados "
    "e diagramas BPMN/UML conforme o template oficial."
)
set_run_font(r, size=9, color=RGBColor(89, 89, 89))

doc.save(path)
print(path)
print("bytes", path.stat().st_size)
